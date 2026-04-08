from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime
from typing import List, Dict, Any
from app.services.reports.raw_data_labels import format_raw_data

BRAND_PURPLE = RGBColor(0x4B, 0x16, 0x4F)
COLOR_SUPERIOR = RGBColor(0x16, 0xa3, 0x4a)
COLOR_NORMAL = RGBColor(0x25, 0x63, 0xeb)
COLOR_BORDERLINE = RGBColor(0xd9, 0x77, 0x06)
COLOR_DEFICIT = RGBColor(0xdc, 0x26, 0x26)

def _classification_color(clasificacion: str) -> RGBColor:
    return {
        'Superior': COLOR_SUPERIOR,
        'Normal': COLOR_NORMAL,
        'Limítrofe': COLOR_BORDERLINE,
        'Deficitario': COLOR_DEFICIT,
    }.get(clasificacion, RGBColor(0, 0, 0))

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def generate_word_report(
    patient: Dict[str, Any],
    sessions: List[Dict[str, Any]],
    plan: Dict[str, Any] | None = None,
    protocol_name: str | None = None,
) -> bytes:
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Nóos — Informe Neuropsicológico', 0)
    title.runs[0].font.color.rgb = BRAND_PURPLE
    doc.add_paragraph(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()

    # Patient section
    h = doc.add_heading('Datos del Paciente', level=1)
    h.runs[0].font.color.rgb = BRAND_PURPLE

    laterality_map = {'diestro': 'Diestro', 'zurdo': 'Zurdo', 'ambidextro': 'Ambidextro'}
    patient_info = [
        ("ID Paciente", patient.get('display_id', patient.get('id', '—'))),
        ("Edad", f"{patient.get('age', '—')} años"),
        ("Escolaridad", f"{patient.get('education_years', '—')} años"),
        ("Lateralidad", laterality_map.get(patient.get('laterality', ''), '—')),
    ]
    if protocol_name:
        patient_info.append(("Protocolo", protocol_name))

    pt = doc.add_table(rows=len(patient_info), cols=2)
    pt.style = 'Table Grid'
    for i, (label, value) in enumerate(patient_info):
        pt.rows[i].cells[0].text = label
        pt.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        pt.rows[i].cells[1].text = value
    doc.add_paragraph()

    # Results section
    h2 = doc.add_heading('Resultados de las Pruebas', level=1)
    h2.runs[0].font.color.rgb = BRAND_PURPLE

    rt = doc.add_table(rows=1 + len(sessions), cols=5)
    rt.style = 'Table Grid'
    headers = ["Prueba", "PE", "Percentil", "Z-Score", "Clasificación"]
    for i, h_text in enumerate(headers):
        cell = rt.rows[0].cells[i]
        cell.text = h_text
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, '4B164F')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_i, s in enumerate(sessions, 1):
        cs = s.get('calculated_scores') or {}
        pe = cs.get('puntuacion_escalar', '—')
        pct = cs.get('percentil', '—')
        z = cs.get('z_score', '—')
        cls_ = cs.get('clasificacion', '—')
        values = [
            s.get('test_type', '—'),
            str(pe),
            f"{pct:.1f}" if isinstance(pct, float) else str(pct),
            f"{z:.2f}" if isinstance(z, float) else str(z),
            cls_,
        ]
        for col_i, val in enumerate(values):
            cell = rt.rows[row_i].cells[col_i]
            cell.text = val
            if col_i == 4:
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
                run.text = val
                run.font.color.rgb = _classification_color(cls_)
                run.bold = True

    # Qualitative notes
    has_notes = any((s.get('qualitative_data') or {}).get('observaciones') for s in sessions)
    if has_notes:
        doc.add_paragraph()
        h3 = doc.add_heading('Observaciones Clínicas', level=1)
        h3.runs[0].font.color.rgb = BRAND_PURPLE
        for s in sessions:
            qd = s.get('qualitative_data') or {}
            obs = qd.get('observaciones', '').strip()
            if obs:
                p = doc.add_paragraph()
                r = p.add_run(f"{s.get('test_type')}: ")
                r.bold = True
                p.add_run(obs)

    # Raw data per test
    doc.add_paragraph()
    h4 = doc.add_heading('Datos Introducidos por Prueba', level=1)
    h4.runs[0].font.color.rgb = BRAND_PURPLE

    for s in sessions:
        test_type = s.get('test_type', '—')
        raw = s.get('raw_data') or {}
        rows_data = format_raw_data(test_type, raw)
        if not rows_data:
            continue

        p = doc.add_paragraph()
        r = p.add_run(test_type)
        r.bold = True
        r.font.color.rgb = BRAND_PURPLE

        rt2 = doc.add_table(rows=len(rows_data), cols=2)
        rt2.style = 'Table Grid'
        for row_i, (lbl, val) in enumerate(rows_data):
            rt2.rows[row_i].cells[0].text = lbl
            if rt2.rows[row_i].cells[0].paragraphs[0].runs:
                rt2.rows[row_i].cells[0].paragraphs[0].runs[0].bold = True
                rt2.rows[row_i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            rt2.rows[row_i].cells[1].text = val
            if rt2.rows[row_i].cells[1].paragraphs[0].runs:
                rt2.rows[row_i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph("Nóos — Triune Neuropsicología · Informe generado automáticamente")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
